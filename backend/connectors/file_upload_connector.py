"""
File Upload Connector for local documents (PDF, DOCX, TXT, MD, JSON, CSV).
Allows users to upload documents directly from their filesystem.
"""

import logging
import json
import csv
import io
from typing import List, Dict, Optional, Any
from pathlib import Path

# PDF parsing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from connectors.base_connector import BaseConnector, ConnectorMetadata

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class FileUploadConnector(BaseConnector):
    """
    Universal file upload connector supporting multiple document formats.

    Supported formats:
    - PDF (.pdf) - Extracts text using PyPDF2 or pdfplumber
    - DOCX (.docx) - Extracts text using python-docx
    - Plain text (.txt, .md) - Direct text reading
    - JSON (.json) - Structured document arrays
    - CSV (.csv) - Tabular data converted to documents
    """

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt', '.md', '.json', '.csv'}

    @classmethod
    def get_metadata(cls) -> ConnectorMetadata:
        """Return connector metadata."""
        return ConnectorMetadata(
            name="File Upload",
            source_type="file_upload",
            description="Upload local documents (PDF, DOCX, TXT, MD, JSON, CSV)",
            version="1.0.0",
            author="RAG Factory",
            category="public",
            supports_incremental_sync=False,  # Files are uploaded once
            supports_rate_limiting=False,  # No external API
            required_config_fields=["files"],
            optional_config_fields=["title_field", "content_field", "max_file_size_mb"],
            default_rate_limit_preset=None
        )

    def __init__(self, config: Dict[str, Any], rate_limit_config: Optional[Dict] = None):
        """
        Initialize file upload connector.

        Args:
            config: Configuration dict with:
                - files (required): List of file data dicts with:
                    - filename: str
                    - content: bytes or str (base64)
                    - mime_type: str (optional)
                - title_field (optional): For JSON/CSV, field name for title
                - content_field (optional): For JSON/CSV, field name for content
                - max_file_size_mb (optional): Max file size in MB (default: 10)
            rate_limit_config: Not used (no external API)
        """
        super().__init__(config, rate_limit_config)

        if 'files' not in config:
            raise ValueError("FileUploadConnector requires 'files' in config")

        self.files = config['files']
        self.title_field = config.get('title_field', 'title')
        self.content_field = config.get('content_field', 'content')
        self.max_file_size_mb = config.get('max_file_size_mb', 10)
        self.max_file_size_bytes = self.max_file_size_mb * 1024 * 1024

        logger.info(f"✓ File Upload connector initialized with {len(self.files)} files")

    def _parse_pdf(self, file_content: bytes, filename: str) -> str:
        """Extract text from PDF file."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 and pdfplumber are required for PDF parsing")

        text_parts = []

        try:
            # Try pdfplumber first (better text extraction)
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                        logger.debug(f"Extracted {len(page_text)} chars from page {page_num}")
        except Exception as e:
            logger.warning(f"pdfplumber failed for {filename}, falling back to PyPDF2: {e}")

            # Fallback to PyPDF2
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            except Exception as e2:
                logger.error(f"PyPDF2 also failed for {filename}: {e2}")
                raise

        full_text = '\n\n'.join(text_parts)
        logger.info(f"✓ Extracted {len(full_text)} characters from PDF: {filename}")
        return full_text

    def _parse_docx(self, file_content: bytes, filename: str) -> str:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX parsing")

        try:
            doc = Document(io.BytesIO(file_content))
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)

            full_text = '\n\n'.join(text_parts)
            logger.info(f"✓ Extracted {len(full_text)} characters from DOCX: {filename}")
            return full_text

        except Exception as e:
            logger.error(f"Failed to parse DOCX {filename}: {e}")
            raise

    def _parse_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from plain text files (TXT, MD)."""
        try:
            # Try UTF-8 first
            text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1
            try:
                text = file_content.decode('latin-1')
                logger.warning(f"Used latin-1 encoding for {filename}")
            except Exception as e:
                logger.error(f"Failed to decode text file {filename}: {e}")
                raise

        logger.info(f"✓ Read {len(text)} characters from text file: {filename}")
        return text

    def _parse_json(self, file_content: bytes, filename: str) -> List[Dict[str, Any]]:
        """Parse JSON file containing document array."""
        try:
            text = file_content.decode('utf-8')
            data = json.loads(text)

            # If it's an array, return it
            if isinstance(data, list):
                logger.info(f"✓ Parsed {len(data)} documents from JSON: {filename}")
                return data

            # If it's a single object, wrap it
            logger.info(f"✓ Parsed 1 document from JSON: {filename}")
            return [data]

        except Exception as e:
            logger.error(f"Failed to parse JSON {filename}: {e}")
            raise

    def _parse_csv(self, file_content: bytes, filename: str) -> List[Dict[str, Any]]:
        """Parse CSV file and convert rows to documents."""
        try:
            text = file_content.decode('utf-8')
            csv_reader = csv.DictReader(io.StringIO(text))
            documents = list(csv_reader)

            logger.info(f"✓ Parsed {len(documents)} rows from CSV: {filename}")
            return documents

        except Exception as e:
            logger.error(f"Failed to parse CSV {filename}: {e}")
            raise

    def fetch_documents(
        self,
        limit: int = 100,
        offset: int = 0,
        since: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """
        Process uploaded files and extract documents.

        Args:
            limit: Maximum number of documents to return
            offset: Number of documents to skip
            since: Not used (file upload doesn't support incremental sync)

        Returns:
            List of documents with id, title, content
        """
        all_documents = []

        for file_idx, file_data in enumerate(self.files):
            filename = file_data.get('filename', f'file_{file_idx}')
            file_extension = Path(filename).suffix.lower()

            # Validate extension
            if file_extension not in self.SUPPORTED_EXTENSIONS:
                logger.warning(f"Skipping unsupported file type: {filename} ({file_extension})")
                continue

            # Get file content
            file_content = file_data.get('content')
            if not file_content:
                logger.warning(f"Skipping file with no content: {filename}")
                continue

            # Convert to bytes if needed
            if isinstance(file_content, str):
                # Assume base64 encoded
                import base64
                try:
                    file_content = base64.b64decode(file_content)
                except Exception as e:
                    logger.error(f"Failed to decode base64 content for {filename}: {e}")
                    continue

            # Check file size
            if len(file_content) > self.max_file_size_bytes:
                logger.warning(
                    f"Skipping file exceeding size limit: {filename} "
                    f"({len(file_content) / 1024 / 1024:.1f}MB > {self.max_file_size_mb}MB)"
                )
                continue

            try:
                # Parse based on file type
                if file_extension == '.pdf':
                    text = self._parse_pdf(file_content, filename)
                    documents = [{
                        'id': filename,
                        'title': filename,
                        'content': text
                    }]

                elif file_extension == '.docx':
                    text = self._parse_docx(file_content, filename)
                    documents = [{
                        'id': filename,
                        'title': filename,
                        'content': text
                    }]

                elif file_extension in {'.txt', '.md'}:
                    text = self._parse_text(file_content, filename)
                    documents = [{
                        'id': filename,
                        'title': filename,
                        'content': text
                    }]

                elif file_extension == '.json':
                    json_docs = self._parse_json(file_content, filename)
                    documents = []
                    for idx, doc in enumerate(json_docs):
                        doc_id = doc.get('id', f'{filename}_{idx}')
                        title = doc.get(self.title_field, filename)
                        content = doc.get(self.content_field, str(doc))
                        documents.append({
                            'id': str(doc_id),
                            'title': title,
                            'content': content
                        })

                elif file_extension == '.csv':
                    csv_rows = self._parse_csv(file_content, filename)
                    documents = []
                    for idx, row in enumerate(csv_rows):
                        doc_id = row.get('id', f'{filename}_row_{idx}')
                        title = row.get(self.title_field, f'Row {idx}')
                        content = row.get(self.content_field, ', '.join(f'{k}: {v}' for k, v in row.items()))
                        documents.append({
                            'id': str(doc_id),
                            'title': title,
                            'content': content
                        })

                else:
                    logger.warning(f"Unsupported file type: {filename}")
                    continue

                # Add metadata
                for doc in documents:
                    doc['metadata'] = {
                        'source': 'file_upload',
                        'filename': filename,
                        'file_type': file_extension,
                        **file_data.get('metadata', {})
                    }

                all_documents.extend(documents)
                logger.info(f"✓ Processed {len(documents)} documents from {filename}")

            except Exception as e:
                logger.error(f"Failed to process file {filename}: {e}", exc_info=True)
                continue

        # Apply offset and limit
        total = len(all_documents)
        result = all_documents[offset:offset + limit]

        logger.info(f"✓ Returning {len(result)} of {total} total documents")
        return result


if __name__ == '__main__':
    """Test the file upload connector"""
    print("=" * 70)
    print("Testing File Upload Connector")
    print("=" * 70)

    # Display connector metadata
    metadata = FileUploadConnector.get_metadata()
    print(f"\n📋 Connector Metadata:")
    print(f"  Name: {metadata.name}")
    print(f"  Type: {metadata.source_type}")
    print(f"  Category: {metadata.category}")
    print(f"  Supported Extensions: {', '.join(FileUploadConnector.SUPPORTED_EXTENSIONS)}")

    # Test with JSON documents
    print("\n📥 Testing with JSON documents...\n")

    test_json = [
        {
            "id": "doc1",
            "title": "Introduction to RAG",
            "content": "Retrieval-Augmented Generation (RAG) is a technique that combines retrieval and generation..."
        },
        {
            "id": "doc2",
            "title": "Vector Databases",
            "content": "Vector databases store embeddings and enable semantic search..."
        }
    ]

    config = {
        'files': [
            {
                'filename': 'test_docs.json',
                'content': json.dumps(test_json).encode('utf-8')
            }
        ]
    }

    connector = FileUploadConnector(config=config)

    # Test connection
    print("🔌 Testing connection...")
    test_result = connector.test_connection()
    print(f"  Status: {'✅ Success' if test_result['success'] else '❌ Failed'}")
    print(f"  Message: {test_result['message']}")

    # Fetch documents
    print("\n📥 Fetching documents...\n")
    docs = connector.fetch_documents()

    if docs:
        print(f"✅ Successfully fetched {len(docs)} documents\n")
        for i, doc in enumerate(docs, 1):
            print(f"{'─' * 70}")
            print(f"Document {i}:")
            print(f"  ID: {doc['id']}")
            print(f"  Title: {doc['title']}")
            print(f"  Content: {doc['content'][:100]}...")
            print(f"  Source: {doc['metadata']['filename']}")
            print()
    else:
        print("❌ No documents retrieved")

    print("=" * 70)
